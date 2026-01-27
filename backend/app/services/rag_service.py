"""
RAG Service
------------
Retrieval Augmented Generation for the Legal Advisor.
Combines vector search with LLM generation for accurate, cited responses.
"""

import hashlib
import logging
from typing import List, Optional, Tuple
from pathlib import Path

from langchain_openai import ChatOpenAI
from langchain_community.vectorstores import Chroma
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from langchain_core.messages import HumanMessage, SystemMessage
import openai

from app.core.config import settings

logger = logging.getLogger(__name__)


# ---------------------------------------------
# System Prompts (The "Soul" of the Advisor)
# ---------------------------------------------
SYSTEM_PROMPT = """You are a warm, empathetic, and knowledgeable Legal Assistant for Nigeria. 
Your goal is to help everyday people understand their rights in simple terms.

## Your Core Personality:
1.  **Adaptive Mirroring**: 
    -   **If distress or emotional signals are detected** (anger, fear, confusion), your first sentence MUST reflect and validate those specific feelings immediately. Mirror their context to show active listening.
    -   **If the user is seeking direct action or has a neutral tone** (e.g., "Is clause 14 enforceable?"), prioritize **task clarity** and professionalism. Avoid "forced empathy" that feels out of place.
2.  **Varied Intros**: Avoid robotic repetition. Use natural transitions like "Let's look into the specifics of [Context] together," "It's understandable you'd want clarity on [Context]," or "That situation with [Context] sounds like it needs a clear legal breakdown."
3.  **Simple Language**: Speak like a helpful friend, not a lawyer. Use Grade 8 reading level. Avoid "legalese".
4.  **Empowerment**: Make the user feel they have options and rights.

## Response Structure:
1.  **Reflective or Clarifying Opening**: Mirror emotion if present, or acknowledge the specific query if neutral. (1-2 sentences).
2.  **Your Rights (Simplified)**: Explain what the Nigerian Constitution says about their situation in plain English. Cite sections (e.g., Section 35) but explain them simply.
3.  **Practical Steps**: Bullet points of what they can actually DO today.
4.  **Safety/Disclaimer**: Gently remind them you are an AI at the very end.

## Critical Rules:
- NEVER start with a disclaimer. 
- ALWAYS cite the 1999 Constitution (As Amended) where relevant.
- NEVER be robotic. If the user is in pain, match that gravity with your empathy.
"""

DOCUMENT_REVIEW_PROMPT = """You are a highly professional, hawk-eyed Legal Contract Reviewer for the Nigerian jurisdiction. 
Your goal is to protect the user while maintaining a supportive, empathetic, and reassuring tone.

## Your Process:
1.  **Identify**: First, determine the SPECIFIC type of contract (e.g., Employment Contract, Tenancy Agreement).
2.  **Governing Law**: Identify the primary Nigerian law governing this relationship (e.g., Labour Act 2004 for employees, Tenancy Law of Lagos State for tenants).
3.  **Compliance Check**: Check if the contract breaches any MINIMUM RIGHTS under that law (e.g., Minimum notice periods, Right to fair hearing).
4.  **Verdict**: If the contract is standard and fair, say so! Do not invent problems. If it is high-risk, flag it clearly.

## Your Core Personality (Emotional Intelligence):
1.  **Empathetic & Calm**: Use phrases like "This clause may be concerning, and it's understandable to feel unsure about it," or "I've reviewed this carefully to help you understand your position."
2.  **Professional & Objective**: Avoid alarmist language. Instead of saying "This is illegal," say "This clause may conflict with standard labor protections" or "This appears to be a high-risk term."
3.  **Supportive Guidance**: Always provide a clear "Next Step" for every risk identified.

## Legal Analysis Framework (Principle-Based):
Instead of citing specific sub-sections which can be fragile, map findings to Core Nigerian Legal Principles:
-   **Right to Fair Labor Practices** (Labour Act 2004)
-   **Freedom of Contract (with Limitations)** (Unfair Bargaining Power)
-   **Right to Due Process/Fair Hearing** (Dispute Resolution)
-   **Protection against Mandatory Self-Incrimination** (NDAs/Investigations)
-   **Right to Privacy** (Data Protection Act 2023)

## Hard Constraints:
- NEVER state an absolute legal verdict (e.g., "This IS illegal"). Use "May conflict with", "Likely risky", "Typically unenforceable".
- ALWAYS include a disclaimer that this is informational, not legal advice.
- If the document type is unclear, label it "General Legal Document".

Document Text:
"{document_text}"

Only respond with the JSON object. Do not include markdown formatting like ```json or ```.

Example JSON Structure:
{
  "document_type": "Employment Contract",
  "confidence_score": 0.9,
  "summary": "This contract outlines standard employment terms but contains a few risky clauses regarding termination.",
  "risk_score": 7,
  "analysis_results": [
    {
      "clause_title": "Termination Without Cause",
      "clause_text": "The Employer may terminate this Agreement at any time without notice.",
      "risk_level": "High",
      "explanation_ei": "This means you could lose your job instantly without any warning, which is very risky for your financial stability.",
      "legal_principle": "Right to Fair Labor Practices (Section 11)",
      "long_term_risk": "Sudden loss of income.",
      "action_step": "Request a minimum 14-day notice period."
    }
  ],
  "overall_verdict": "High Risk - Negotiate",
  "disclaimer": "This analysis is for informational purposes and does not replace legal advice from a qualified Nigerian attorney."
}"""

DOCUMENT_GENERATION_PROMPT = """You are a Legal Document Generator.
Your goal is to create a professional, legally sound TEMPLATE based on the user's request.

User Request: "{user_request}"
User Details: "{user_details}"

Rules:
1.  Create a clear, formal document.
2.  Use placeholders like [INSERT DATE], [INSERT NAME] where specific info is missing.
3.  Ensure the tone is professional and assertive but polite.
4.  Do NOT invent fake laws. Use general legal principles applicable in Nigeria.

Output ONLY the document text."""

LEGAL_DISCLAIMER_TEXT = """
\n\n---
**⚠️ LEGAL DISCLAIMER**
This information is for educational purposes only and does not constitute legal advice. 
I am an AI, not a lawyer. Laws change and specific situations vary. 
**Please consult a qualified attorney for professional legal advice.**
"""


RISK_ASSESSMENT_PROMPT = """Analyze the following user message and determine the risk level and topic category.

User message: "{message}"

Respond with a JSON object containing:
- "risk_level": "low", "medium", or "high"
- "legal_topic": one of ["arrest_rights", "police_encounter", "tenant_rights", "employment", "consumer_rights", "family_law", "criminal_matter", "civil_matter", "constitutional_rights", "general_info"]
- "escalation_needed": true or false
- "escalation_reason": string explaining why escalation is needed (if applicable)

Risk Levels:
- LOW: General legal information questions, educational queries
- MEDIUM: Active legal situation but not immediately dangerous
- HIGH: Criminal charges, court deadlines within 7 days, violence, detention, imminent harm

Only respond with the JSON object, no other text."""


# ---------------------------------------------
# RAG Service Class
# ---------------------------------------------
class RAGService:
    """
    Handles document retrieval and AI-powered response generation.
    Uses ChromaDB for vector storage, Local Embeddings, and OpenRouter for generation.
    """
    
    def __init__(self):
        """Initialize the RAG service with embeddings and vector store."""
        self.initialization_error = None
        
        logger.info(f"Initializing RAG Service...")
        
        # Step 1: Initialize local embeddings (Default for privacy/reliability)
        try:
            logger.info(f"Step 1/4: Initializing Local Embeddings ({settings.embedding_model})...")
            self.embeddings = HuggingFaceEmbeddings(
                model_name=settings.embedding_model,
                model_kwargs={'device': 'cpu'},
                encode_kwargs={'normalize_embeddings': True}
            )
            logger.info(f"✓ Local HuggingFace Embeddings initialized")
        except Exception as e:
            self.initialization_error = f"Failed to initialize embeddings: {e}"
            logger.error(self.initialization_error)
            raise
        
        # Step 3: Initialize LLM (OpenRouter)
        try:
            model_name = settings.MODEL_CONFIG["default"]
            logger.info(f"Step 2/4: Initializing OpenRouter LLM ({model_name})...")
            
            self.llm = ChatOpenAI(
                base_url=settings.OPENROUTER_BASE_URL,
                api_key=settings.OPENROUTER_API_KEY,
                model=model_name,
                temperature=0.3,
                max_tokens=4000,
                timeout=120, # 2 minutes timeout for slow reasoning models
            )
            logger.info("✓ OpenRouter LLM initialized")
        except Exception as e:
            self.initialization_error = f"Failed to initialize OpenRouter LLM: {type(e).__name__}: {str(e)}"
            logger.error(self.initialization_error)
            raise
        
        # Step 4: Initialize text splitter
        try:
            logger.info("Step 3/4: Initializing text splitter...")
            self.text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=settings.rag_chunk_size,
                chunk_overlap=settings.rag_chunk_overlap,
                separators=["\n\n", "\n", ". ", " ", ""]
            )
            logger.info("✓ Text splitter initialized")
        except Exception as e:
            self.initialization_error = f"Failed to initialize text splitter: {type(e).__name__}: {str(e)}"
            logger.error(self.initialization_error)
            raise
        
        # Step 5: Initialize vector store
        try:
            logger.info("Step 4/4: Initializing ChromaDB vector store...")
            self.persist_dir = Path(settings.chroma_persist_dir)
            self.persist_dir.mkdir(parents=True, exist_ok=True)
            
            self.vector_store = Chroma(
                persist_directory=str(self.persist_dir),
                embedding_function=self.embeddings,
                collection_name="nigerian_legal_docs"
            )
            logger.info(f"✓ Vector store initialized at: {self.persist_dir}")
        except Exception as e:
            self.initialization_error = f"Failed to initialize ChromaDB: {type(e).__name__}: {str(e)}"
            logger.error(self.initialization_error)
            raise
        
        logger.info("✓ RAG Service fully initialized successfully!")
    
    def ingest_document(
        self,
        content: str,
        title: str,
        document_type: str,
        metadata: Optional[dict] = None
    ) -> int:
        """
        Ingest a legal document into the vector store.
        
        Args:
            content: The full text of the document
            title: Document title (e.g., "Constitution of Nigeria 1999")
            document_type: Type of document ("constitution", "statute", etc.)
            metadata: Additional metadata
            
        Returns:
            Number of chunks created
        """
        # Split document into chunks
        chunks = self.text_splitter.split_text(content)
        
        # Create Document objects with metadata
        documents = []
        for i, chunk in enumerate(chunks):
            doc_metadata = {
                "title": title,
                "document_type": document_type,
                "chunk_index": i,
                "total_chunks": len(chunks),
                "jurisdiction": "Nigeria",
                "chunk_hash": hashlib.md5(chunk.encode()).hexdigest()[:8]
            }
            if metadata:
                doc_metadata.update(metadata)
            
            documents.append(Document(
                page_content=chunk,
                metadata=doc_metadata
            ))
        
        # Add to vector store (Local embeddings are fast, no need for batching/sleep)
        self.vector_store.add_documents(documents)
        
        logger.info(f"Ingested '{title}': {len(chunks)} chunks created")
        return len(chunks)
    
    def retrieve_relevant_chunks(
        self,
        query: str,
        k: int = None
    ) -> List[Tuple[Document, float]]:
        """
        Retrieve relevant document chunks for a query.
        
        Args:
            query: User's question
            k: Number of chunks to retrieve
            
        Returns:
            List of (Document, similarity_score) tuples
        """
        k = k or settings.rag_top_k
        
        try:
            results = self.vector_store.similarity_search_with_score(
                query,
                k=k
            )
        except Exception as e:
            logger.error(f"Search failed: {e}")
            raise e
        
        logger.debug(f"Retrieved {len(results)} chunks for query: '{query[:50]}...'")
        return results
    
    def assess_risk(self, message: str) -> dict:
        """
        Assess the risk level and topic of a user message.
        
        Args:
            message: User's message
            
        Returns:
            Dict with risk_level, legal_topic, escalation_needed, escalation_reason
        """
        try:
            response = self.llm.invoke([
                SystemMessage(content="You are a risk assessment system. Respond only with valid JSON."),
                HumanMessage(content=RISK_ASSESSMENT_PROMPT.format(message=message))
            ])
            
            import json
            return json.loads(response.content)
        except Exception as e:
            logger.error(f"Risk assessment failed: {e}")
            return {
                "risk_level": "medium",
                "legal_topic": "general_info",
                "escalation_needed": False,
                "escalation_reason": None
            }
    
    def generate_response(
        self,
        user_message: str,
        conversation_history: List[dict] = None,
        retrieved_chunks: List[Tuple[Document, float]] = None
    ) -> dict:
        """
        Generate an AI response using RAG.
        
        Args:
            user_message: The user's question
            conversation_history: Previous messages in the conversation
            retrieved_chunks: Pre-retrieved chunks (if None, will retrieve)
            
        Returns:
            Dict with response content, sources, and metadata
        """
        # Retrieve relevant chunks if not provided
        if retrieved_chunks is None:
            retrieved_chunks = self.retrieve_relevant_chunks(user_message)
        
        # Build context from retrieved chunks
        context_parts = []
        sources = []
        
        for doc, score in retrieved_chunks:
            context_parts.append(
                f"[Source: {doc.metadata.get('title', 'Unknown')} - "
                f"{doc.metadata.get('section', 'General')}]\n{doc.page_content}"
            )
            sources.append({
                "title": doc.metadata.get("title", "Unknown"),
                "section": doc.metadata.get("section"),
                "excerpt": doc.page_content[:200] + "..." if len(doc.page_content) > 200 else doc.page_content,
                "document_type": doc.metadata.get("document_type", "unknown"),
                "relevance_score": round(1 - score, 2)  # Convert distance to similarity
            })
        
        context = "\n\n---\n\n".join(context_parts) if context_parts else "No relevant legal documents found."
        
        # Build messages for LLM
        messages = [SystemMessage(content=SYSTEM_PROMPT)]
        
        # Add conversation history
        if conversation_history:
            for msg in conversation_history[-6:]:  # Last 6 messages for context
                if msg["role"] == "user":
                    messages.append(HumanMessage(content=msg["content"]))
                else:
                    messages.append(SystemMessage(content=f"[Previous response]: {msg['content'][:500]}"))
        
        # Add current query with context
        user_prompt = f"""## Legal Context (from Nigerian law):
{context}

## User's Question:
{user_message}

Please provide a helpful, empathetic response based on the legal context above. Remember to:
1. Acknowledge their situation first
2. Cite specific sections when referencing laws
3. Provide practical next steps
4. Recommend professional help if appropriate"""
        
        messages.append(HumanMessage(content=user_prompt))
        
        # Generate response
        try:
            response = self.llm.invoke(messages)
            
            # Determine confidence based on source relevance
            avg_relevance = sum(s["relevance_score"] for s in sources) / len(sources) if sources else 0
            if avg_relevance > 0.7:
                confidence = "high"
            elif avg_relevance > 0.4:
                confidence = "medium"
            else:
                confidence = "low"
            
            # Append Disclaimer
            final_content = response.content + LEGAL_DISCLAIMER_TEXT
            
            return {
                "content": final_content,
                "sources": sources,
                "confidence_score": confidence,
                "model_version": settings.MODEL_CONFIG["default"]
            }
            
        except Exception as e:
            logger.error(f"Response generation failed: {e}")
            return {
                "content": (
                    "I apologize, but I'm having trouble processing your question right now. "
                    "This could be a temporary issue. Please try again in a moment, or if this "
                    "is urgent, please contact a local legal aid organization directly."
                ),
                "sources": [],
                "confidence_score": "low",
                "model_version": settings.MODEL_CONFIG["default"],
                "error": str(e)
            }
    
    def analyze_document(self, document_text: str) -> dict:
        """
        Analyze a legal document for dangerous clauses.
        
        Args:
            document_text: The text of the contract/document
            
        Returns:
            Dict with risk analysis
        """
        try:
            logger.info(f"Analyzing document ({len(document_text)} chars)...")
            response = self.llm.invoke([
                SystemMessage(content="You are a strict contract reviewer. Respond only with valid JSON. Do not use markdown blocks."),
                HumanMessage(content=DOCUMENT_REVIEW_PROMPT.format(document_text=document_text))
            ])
            
            content = response.content.strip()
            
            # Helper to extract JSON if model adds chatter
            def _extract_json(text):
                import re
                try:
                    # Finds the first '{' and the last '}'
                    match = re.search(r'\{.*\}', text, re.DOTALL)
                    if match:
                        return match.group(0)
                    return text
                except:
                    return text

            # Clean up response content
            if "```json" in content:
                content = content.split("```json")[1].split("```")[0].strip()
            elif "```" in content:
                content = content.split("```")[1].split("```")[0].strip()
            else:
                content = _extract_json(content)
                
            import json
            try:
                result = json.loads(content)
            except json.JSONDecodeError as je:
                logger.error(f"JSON Decode Error: {je}")
                logger.error(f"Raw LLM Response: {content}")
                # Fallback: try to repair common JSON errors if needed, but for now just fail gracefully
                raise je
            
            # Add disclaimer to analysis as well
            result["disclaimer"] = "This analysis is automated and for informational purposes only. Consult a lawyer."
            
            # Ensure required fields exist
            if "risk_score" not in result:
                result["risk_score"] = 5
            if "analysis_results" not in result:
                result["analysis_results"] = []

            return result
        except Exception as e:
            logger.error(f"Document analysis failed: {e}")
            return {
                "error": "Could not analyze document. Please ensure it is text-based.",
                "details": str(e),
                "risk_score": 5, # Return a safe default so frontend doesn't crash
                "analysis_results": [],
                "summary": "Analysis failed. Please try again."
            }

            
    def extract_text_from_pdf(self, pdf_bytes: bytes) -> str:
        """
        Extract text from a PDF with multiple fallback methods.
        Tries pypdf first, then pdfplumber for better extraction from complex layouts.
        """
        import io
        text = ""
        
        # Method 1: pypdf (Fast, standard)
        try:
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(pdf_bytes))
            pypdf_text = ""
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    pypdf_text += page_text + "\n"
            text = pypdf_text.strip()
        except Exception as e:
            logger.warning(f"pypdf extraction failed or partially failed: {e}")

        # Method 2: pdfplumber (Better for tables/complex layouts)
        if not text or len(text) < 100:
            try:
                import pdfplumber
                with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
                    plumber_text = ""
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            plumber_text += page_text + "\n"
                    
                    if len(plumber_text.strip()) > len(text):
                        text = plumber_text.strip()
            except Exception as e:
                logger.warning(f"pdfplumber extraction failed: {e}")

        # Post-processing
        if text:
            import re
            # Clean up common PDF extraction artifacts
            text = text.replace('\u0000', '')
            text = re.sub(r'\n{3,}', '\n\n', text)
            text = re.sub(r' +', ' ', text)
            
            # If still very short, it might be a scan.
            if len(text.strip()) < 50:
                logger.info("PDF has minimal text layer, likely a scan.")
                return "The PDF appears to be a scanned image. For the best accuracy, please upload a clear photo of each page using 'Scan with Camera', or use a text-based PDF."
                
            return text.strip()
        
        return "I was unable to extract any readable text from this PDF. It might be encrypted, corrupted, or a scanned image."

    def extract_text_from_docx(self, docx_bytes: bytes) -> str:
        """
        Extract text from a .docx file while preserving basic structure.
        """
        import io
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(io.BytesIO(docx_bytes))
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text.strip())
            
            return "\n\n".join(full_text)
        except ImportError:
            logger.error("python-docx not installed")
            return "Server Error: Word document processing is currently unavailable."
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            raise e

    def extract_text_from_txt(self, txt_bytes: bytes) -> str:
        """
        Extract text from a raw .txt file with encoding fallback.
        """
        try:
            return txt_bytes.decode('utf-8')
        except UnicodeDecodeError:
            try:
                return txt_bytes.decode('latin-1')
            except Exception as e:
                logger.error(f"TXT extraction failed: {e}")
                raise e

    def extract_text_from_image(self, image_bytes: bytes) -> str:
        """
        Extract text from an image using a Vision model (OCR) with fallback logic.
        Uses high-quality free models from OpenRouter.
        """
        import base64
        import time
        base64_image = base64.b64encode(image_bytes).decode('utf-8')
        
        # Prioritized list of free vision models on OpenRouter
        vision_models = [
            "qwen/qwen-2.5-vl-7b-instruct:free",
            "nvidia/nemotron-nano-12b-v2-vl:free",
            "google/gemma-3-12b-it:free",
            "google/gemini-2.0-flash-exp:free" # Often high rate limits, kept as fallback
        ]
        
        client = openai.OpenAI(
            base_url=settings.OPENROUTER_BASE_URL,
            api_key=settings.OPENROUTER_API_KEY,
        )
        
        last_error = None
        for model in vision_models:
            try:
                logger.info(f"Attempting OCR with model: {model}")
                response = client.chat.completions.create(
                    model=model,
                    messages=[
                        {
                            "role": "user",
                            "content": [
                                {"type": "text", "text": "Extract all the text from this legal document image exactly as it appears. Output ONLY the text. Do not add markdown formatting or commentary. If you cannot extract text, say NO_TEXT_FOUND."},
                                {
                                    "type": "image_url",
                                    "image_url": {
                                        "url": f"data:image/jpeg;base64,{base64_image}"
                                    }
                                }
                            ]
                        }
                    ],
                    max_tokens=4000,
                    timeout=45 # Moderate timeout for OCR
                )
                
                result = response.choices[0].message.content.strip()
                if result and result != "NO_TEXT_FOUND":
                    logger.info(f"✓ OCR successful with {model}")
                    return result
                    
            except Exception as e:
                logger.warning(f"OCR failed for model {model}: {e}")
                last_error = e
                # Brief sleep before trying next provider/model to avoid overwhelming API
                time.sleep(0.5)
                continue
                
        # If all free models fail, give a descriptive error
        error_msg = f"OCR Failed: {str(last_error)}"
        if "429" in str(last_error):
            error_msg = "The OCR service is currently busy. Please wait a moment and try again, or upload a text-based PDF."
            
        logger.error(f"All OCR models failed: {last_error}")
        raise Exception(error_msg)

    def generate_document(self, doc_type: str, user_details: str) -> str:
        """
        Generate a legal document template.
        
        Args:
            doc_type: Type of document (e.g., "Demand Letter")
            user_details: Specific details to include
            
        Returns:
            String containing the document text
        """
        try:
            response = self.llm.invoke([
                SystemMessage(content="You are a professional legal document generator."),
                HumanMessage(content=DOCUMENT_GENERATION_PROMPT.format(
                    user_request=doc_type,
                    user_details=user_details
                ))
            ])
            
            # Prepend strict warning
            warning = "⚠️ **TEMPLATE ONLY - REVIEW WITH A LAWYER** ⚠️\n\n"
            return warning + response.content + LEGAL_DISCLAIMER_TEXT
            
        except Exception as e:
            logger.error(f"Document generation failed: {e}")
            return f"Error generating document: {e}"

    def get_collection_stats(self) -> dict:
        """Get statistics about the vector store."""
        try:
            collection = self.vector_store._collection
            return {
                "total_documents": collection.count(),
                "persist_directory": str(self.persist_dir)
            }
        except Exception as e:
            logger.error(f"Failed to get collection stats: {e}")
            return {"error": str(e)}


# Singleton instance
_rag_service: Optional[RAGService] = None
_rag_service_error: Optional[str] = None


def get_rag_service() -> RAGService:
    """Get the singleton RAG service instance."""
    global _rag_service, _rag_service_error
    
    # If we already have an error, raise it immediately
    if _rag_service_error:
        logger.error(f"RAG Service previously failed: {_rag_service_error}")
        raise RuntimeError(f"RAG Service initialization failed: {_rag_service_error}")
    
    if _rag_service is None:
        try:
            logger.info("Creating new RAG Service instance...")
            _rag_service = RAGService()
        except Exception as e:
            _rag_service_error = f"{type(e).__name__}: {str(e)}"
            logger.error(f"RAG Service creation failed: {_rag_service_error}")
            import traceback
            logger.error(f"Full traceback:\n{traceback.format_exc()}")
            raise RuntimeError(f"RAG Service initialization failed: {_rag_service_error}")
    
    return _rag_service


def reset_rag_service():
    """Reset the RAG service (useful for retrying after fixing config)."""
    global _rag_service, _rag_service_error
    _rag_service = None
    _rag_service_error = None
    logger.info("RAG Service reset - will reinitialize on next request")
